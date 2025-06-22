import streamlit as st
import os
import pandas as pd
from datetime import datetime
import base64
from PIL import Image
import docx
import warnings

# Suppress warnings
warnings.filterwarnings("ignore")

# --- Hardcoded credentials ---
USERNAME = "Compiler Design"
PASSWORD = "cse331"

# --- Initialize session state ---
if 'logged_in' not in st.session_state:
    st.session_state['logged_in'] = False
if 'current_slide' not in st.session_state:
    st.session_state['current_slide'] = 0

# --- Utility Functions ---
def login():
    st.title("🔐 Login")
    username = st.text_input("Username")
    password = st.text_input("Password", type="password")
    if st.button("Login"):
        if username == USERNAME and password == PASSWORD:
            st.session_state['logged_in'] = True
            st.session_state['username'] = username
            st.success("Logged in successfully!")
            st.rerun()
        else:
            st.error("Invalid username or password")
    show_footer()

def logout():
    st.session_state['logged_in'] = False
    st.rerun()

def show_navbar():
    with st.sidebar:
        st.title("Course Content Helper")
        st.write("---")
        nav = st.radio(
            "Navigate",
            ("Home", "Course Content", "Contact"),
            label_visibility="collapsed"
        )
        st.write("---")
        if st.button("Logout", use_container_width=True):
            logout()
    return nav

def show_footer():
    st.markdown("---")
    current_year = datetime.now().year
    footer = f"""
    <div style="text-align: center; padding: 10px; color: gray; font-size: 14px;">
        <p>© {current_year} Course Content Helper. All rights reserved.</p>
        <p>Version 1.0.0 | <a href="#" style="color: gray;">Privacy Policy</a> | <a href="#" style="color: gray;">Terms of Service</a></p>
    </div>
    """
    st.markdown(footer, unsafe_allow_html=True)

def show_home():
    st.markdown(
        """
        <div style="text-align:center; padding: 20px;">
            <h1>🏠 Welcome to the Course Content Helper</h1>
            <img src="https://images.unsplash.com/photo-1504384308090-c894fdcc538d?auto=format&fit=crop&w=800&q=80" alt="Home Image" style="width:70%; border-radius:15px;"/>
            <p style="font-size:18px; margin-top:20px;">
                Use the navigation to browse course materials and learn at your own pace.<br>
                Happy Learning! 📚
            </p>
        </div>
        """, unsafe_allow_html=True
    )
    show_footer()

def display_pdf(file_path):
    with open(file_path, "rb") as f:
        base64_pdf = base64.b64encode(f.read()).decode('utf-8')
    
    # Display with object tag instead of iframe
    pdf_display = f'''
    <object data="data:application/pdf;base64,{base64_pdf}" 
            type="application/pdf" 
            width="120%" 
            height="1000px">
    </object>
    '''
    st.markdown(pdf_display, unsafe_allow_html=True)


def display_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    st.text_area("Document Content", "\n".join(full_text), height=800, label_visibility="collapsed")
    
    # Display tables if any
    for i, table in enumerate(doc.tables):
        st.subheader(f"Table {i+1}")
        st.dataframe(pd.DataFrame([[cell.text for cell in row.cells] for row in table.rows]), use_container_width=True)

def show_course_content():
    st.title("📁 Course Materials")
    
    content_dir = "content"
    if not os.path.isdir(content_dir):
        st.error(f"The folder '{content_dir}' does not exist. Please create it and add your course files.")
        show_footer()
        return

    files = [f for f in os.listdir(content_dir) if f.lower().endswith(('.pdf', '.pptx', '.docx'))]
    
    if not files:
        st.info("No supported files found (PDF, PPTX, DOCX) in the content folder.")
        show_footer()
        return

    selected_file = st.selectbox("Select a file to view or download", files)
    file_path = os.path.join(content_dir, selected_file)
    
    tab1, tab2 = st.tabs(["View", "Download"])
    
    with tab1:
        st.subheader(f"Viewing: {selected_file}")
        
        try:
            if selected_file.lower().endswith(".pdf"):
                display_pdf(file_path)
            elif selected_file.lower().endswith(".docx"):
                display_docx(file_path)
            else:
                st.warning("This file type is not supported for preview")
        except Exception as e:
            st.error(f"Error displaying file: {str(e)}")

    with tab2:
        st.subheader("Download Options")
        with open(file_path, "rb") as f:
            st.download_button(
                label="⬇️ Download File",
                data=f,
                file_name=selected_file,
                mime="application/octet-stream",
                use_container_width=True
            )
    
    show_footer()

def show_contact():
    st.title("📬 Contact Us")
    st.markdown(
        """
        <div style="font-size:18px; line-height:1.6;">
            <p><strong>Name:</strong> Azaz Ahamed</p>
            <p><strong>Email:</strong> ahamed15-6216@s.diu.edu.bd</p>
            <p><strong>Phone:</strong> +880 1931 707075</p>
            <p><strong>Address:</strong> Daffodil International University, Dhaka, Bangladesh</p>
        </div>
        """,
        unsafe_allow_html=True
    )
    show_footer()

# --- Main App ---
if not st.session_state['logged_in']:
    login()
else:
    page = show_navbar()

    if page == "Home":
        show_home()
    elif page == "Course Content":
        show_course_content()
    elif page == "Contact":
        show_contact()